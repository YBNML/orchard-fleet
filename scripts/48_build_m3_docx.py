#!/usr/bin/env python3
"""M3 보고서 md → docx 변환기

    python3 scripts/48_build_m3_docx.py

docs/design/M3_localization_report.md 를 한국어 보고서
docs/design/M3_로컬리제이션_실험보고서.docx 로 변환한다.
지원: 제목(#~####), 표, 이미지(../figures/...), 코드블록, 굵게, 인용, 목록, 구분선.
"""
import pathlib
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
SRC = ROOT / "docs/design/M3_localization_report.md"
OUT = ROOT / "docs/design/M3_로컬리제이션_실험보고서.docx"

INK = RGBColor(0x24, 0x1F, 0x19)
LEAF = RGBColor(0x3A, 0x52, 0x32)
GOLD = RGBColor(0x8A, 0x5D, 0x12)


def set_kr(style, size, bold=False, color=None, mono=False):
    f = style.font
    f.name = "Consolas" if mono else "맑은 고딕"
    f.size = Pt(size)
    f.bold = bold
    if color is not None:
        f.color.rgb = color
    style.element.rPr.rFonts.set(qn("w:eastAsia"),
                                 "Consolas" if mono else "맑은 고딕")


def add_runs(par, text, mono=False):
    """**굵게**·`코드` 인라인 처리."""
    for tok in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = LEAF
        else:
            r = par.add_run(tok)
        if mono:
            r.font.name = "Consolas"
            r.font.size = Pt(9)


def main():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    st = doc.styles["Normal"]
    set_kr(st, 10.5, color=INK)
    for name, size, bold, color in (("Heading 1", 17, True, LEAF),
                                    ("Heading 2", 14, True, INK),
                                    ("Heading 3", 12, True, GOLD),
                                    ("Heading 4", 11, True, INK)):
        set_kr(doc.styles[name], size, bold, color)

    i, n = 0, len(lines)
    while i < n:
        ln = lines[i]
        s = ln.strip()
        if not s:
            i += 1
            continue
        # 코드블록
        if s.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            add_runs(p, "")
            for j, code in enumerate(buf):
                r = p.add_run(code + ("\n" if j < len(buf) - 1 else ""))
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            p.paragraph_format.left_indent = Cm(0.6)
            continue
        # 표
        if s.startswith("|") and i + 1 < n and set(lines[i + 1].strip()) <= set("|-: "):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            rows.pop(1)                          # 구분행 제거
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Light Grid Accent 3"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci >= len(t.rows[ri].cells):
                        continue
                    c = t.rows[ri].cells[ci]
                    c.text = ""
                    p = c.paragraphs[0]
                    add_runs(p, cell)
                    for r in p.runs:
                        r.font.size = Pt(9.5)
                        if ri == 0:
                            r.bold = True
            doc.add_paragraph()
            continue
        # 이미지
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", s)
        if m:
            img = (SRC.parent / m.group(2)).resolve()
            if img.exists():
                doc.add_picture(str(img), width=Cm(15.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(m.group(1))
                r.font.size = Pt(9)
                r.font.color.rgb = GOLD
            i += 1
            continue
        # 구분선
        if s in ("---", "***"):
            i += 1
            continue
        # 제목
        m = re.match(r"(#{1,4})\s+(.*)", s)
        if m:
            doc.add_heading("", level=len(m.group(1)))
            add_runs(doc.paragraphs[-1], m.group(2))
            i += 1
            continue
        # 인용
        if s.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip("> "))
                i += 1
            p = doc.add_paragraph()
            add_runs(p, " ".join(buf))
            p.paragraph_format.left_indent = Cm(0.8)
            for r in p.runs:
                r.font.color.rgb = LEAF
                r.italic = True
            continue
        # 목록
        m = re.match(r"[-*]\s+(.*)", s)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(1))
            i += 1
            continue
        m = re.match(r"(\d+)\.\s+(.*)", s)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue
        # 일반 문단 (연속 줄 병합)
        buf = [s]
        i += 1
        while i < n and lines[i].strip() and not re.match(
                r"(#{1,4}\s|\||>|[-*]\s|\d+\.\s|```|!\[|---)", lines[i].strip()):
            buf.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(buf))

    doc.save(OUT)
    print(f"저장: {OUT} ({OUT.stat().st_size:,} B)")


if __name__ == "__main__":
    main()
