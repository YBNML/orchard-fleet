#!/usr/bin/env python3
"""설계서 워드(.docx) 빌드

docs/superpowers/specs/2026-08-01-orchard-fleet-3tier-design.md (정본)
  → pandoc → 후처리(한글 폰트·표 테두리) → docs/design/통합관제_설계서_v1.docx
"""
from pathlib import Path

import pypandoc
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

import sys
_a = sys.argv[1:]
MD = Path(_a[0]) if _a else Path("docs/superpowers/specs/2026-08-01-orchard-fleet-3tier-design.md")
OUT_DIR = Path("docs/design")
OUT_DIR.mkdir(parents=True, exist_ok=True)
TMP = OUT_DIR / "_pandoc_raw.docx"
OUT = OUT_DIR / (_a[1] if len(_a) > 1 else "통합관제_설계서_v1.docx")

# ── 1. pandoc 변환 ──────────────────────────────────────────────────────────
pypandoc.convert_file(
    str(MD), "docx", format="markdown+pipe_tables+link_attributes",
    outputfile=str(TMP),
    extra_args=[f"--resource-path={MD.parent}", "--toc", "--toc-depth=2",
                "--metadata", "lang=ko-KR",
                "--metadata", "toc-title=목차"],
)

# ── 2. 후처리 ───────────────────────────────────────────────────────────────
doc = Document(str(TMP))

EAST = "맑은 고딕"      # Windows Word 기본 한글 폰트 (없으면 Noto 폴백)
LATIN = "Malgun Gothic"

def set_style_font(style, size=None, east=EAST, latin=LATIN):
    try:
        style.font.name = latin
        if size:
            style.font.size = Pt(size)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), east)
    except Exception:
        pass

for name, size in [("Normal", 10), ("Body Text", 10), ("Title", 22),
                   ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5),
                   ("Heading 4", 10.5), ("Compact", 10), ("Table Caption", 9),
                   ("Image Caption", 9), ("Caption", 9), ("TOC Heading", 14)]:
    try:
        set_style_font(doc.styles[name], size)
    except KeyError:
        pass

# 표 전체: 테두리 + 셀 폰트 축소
def set_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = tblPr.makeelement(qn("w:tblBorders"), {})
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = borders.makeelement(qn(f"w:{edge}"), {})
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "8aa0b8")

for tbl in doc.tables:
    set_table_borders(tbl)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

doc.save(str(OUT))
TMP.unlink()
print(f"생성: {OUT}  ({OUT.stat().st_size//1024} KB)")
